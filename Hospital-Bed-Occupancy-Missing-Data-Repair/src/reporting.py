"""Generate a polished, evidence-based project report in DOCX/PDF form."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


class ProjectReportBuilder:
    """Build a concise executive report using a consistent business-brief visual system."""

    def __init__(self, reports_dir: Path, figures_dir: Path) -> None:
        self.reports_dir = reports_dir
        self.figures_dir = figures_dir

    def build(self, payload: dict[str, Any]) -> Path:
        """Create the editable source report; PDF conversion is handled by the delivery workflow."""
        document = Document()
        self._set_document_tokens(document)
        self._add_header_footer(document)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(4)
        run = title.add_run("Hospital Bed Occupancy Missing-Data Repair")
        self._font(run, "Calibri", 24, "0B2545", bold=True)
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(18)
        run = subtitle.add_run("Advanced Imputation Benchmark and Occupancy Risk Modeling")
        self._font(run, "Calibri", 12, "5D6D7E")
        self._paragraph(document, "Executive summary", "A synthetic 17,520-record hospital census was used to test 12 missing-data repair approaches across MCAR, MAR, and MNAR patterns. The system evaluates reconstruction fidelity, distribution drift, execution time, and downstream high-occupancy classification performance, then selects a strategy per feature.", level=1)
        self._add_key_table(document, payload)
        sections = [
            ("Problem Statement", "Hospital operations data often has incomplete staffing, admissions, weather, and length-of-stay fields. Replacing all gaps with one generic statistic can hide meaningful operational signals and lead to unreliable bed-capacity decisions."),
            ("Objectives", "Compare multiple imputation strategies; choose the most suitable method per feature; quantify downstream model impact; and provide a deployable GUI for repair, inspection, training, prediction, and export."),
            ("Dataset and Missingness Design", "The project uses daily departmental census records from 16 synthetic hospitals across 365 days. MCAR gaps are applied randomly; MAR gaps depend on observed department, city, or weather; MNAR gaps are concentrated in high staffing and high COVID census values. The target occupancy rate is intentionally never masked."),
            ("Data Preparation and Feature Engineering", "The workflow removes duplicates, converts types, visualizes missingness, flags anomalies through IQR and Isolation Forest, imputes selected fields, normalizes exported features, and creates capacity, staffing, calendar, interaction, polynomial, lag, rolling mean, rolling median, and moving-average variables. Temporal features only use prior observations."),
            ("Imputation Comparison", f"The benchmark evaluates Mean, Median, Mode, Forward Fill, Backward Fill, Linear, Polynomial, and Spline interpolation, KNN, MICE, Regression, and Random Forest imputation. The overall top method by mean reconstruction RMSE was {payload['best_imputation_method']}; final selection remains feature-specific rather than global."),
            ("Model Performance", f"{payload['best_model']} was selected using macro-F1 cross-validation on the training period and evaluated on a held-out future period. Holdout accuracy was {payload['accuracy']:.3f}; macro F1 was {payload['f1_macro']:.3f}. The model predicts Low, Moderate, and High occupancy risk classes."),
            ("Business Insights", "Occupancy planning is most sensitive to admissions flow, staffing intensity, emergency load, seasonal patterns, and lagged demand. The feature-specific repair policy preserves these relationships more faithfully than a single default imputer, supporting staffing and bed-allocation decisions."),
            ("Limitations and Future Work", "The data is synthetic and cannot establish clinical causality. Future work should validate against governed real-world data, calibrate thresholds to local policy, use prospective monitoring for drift, add LightGBM/CatBoost where approved, and expose formal audit logging for production deployment."),
        ]
        for heading, body in sections:
            self._paragraph(document, heading, body, level=1)
            if heading == "Dataset and Missingness Design":
                self._image(document, "missing_value_heatmap.png", "Figure 1. Missingness pattern sample")
            elif heading == "Data Preparation and Feature Engineering":
                self._image(document, "monthly_seasonal_holiday_trends.png", "Figure 2. Capacity and seasonal occupancy trends")
            elif heading == "Model Performance":
                self._image(document, "confusion_matrix.png", "Figure 3. Holdout classification performance")
            elif heading == "Business Insights":
                self._image(document, "permutation_importance.png", "Figure 4. Feature importance for operational planning")
        self._add_recommendations_table(document, payload)
        output = self.reports_dir / "Project_Report.docx"
        document.save(output)
        return output

    def _set_document_tokens(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = section.bottom_margin = Inches(0.85)
        section.left_margin = section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.08
        for style_name, size, color in [("Heading 1", 15, "2E74B5"), ("Heading 2", 12, "1F4D78")]:
            style = styles[style_name]
            style.font.name = "Calibri"; style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(13); style.paragraph_format.space_after = Pt(6)

    def _add_header_footer(self, document: Document) -> None:
        section = document.sections[0]
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("Hospital Census Analytics | Project Report")
        self._font(run, "Calibri", 8.5, "7F8C8D")
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Synthetic data only | Portfolio capstone")
        self._font(run, "Calibri", 8.5, "7F8C8D")

    def _paragraph(self, document: Document, title: str, body: str, level: int = 0) -> None:
        if level:
            document.add_heading(title, level=level)
        paragraph = document.add_paragraph(body)
        paragraph.paragraph_format.space_after = Pt(7)

    def _add_key_table(self, document: Document, payload: dict[str, Any]) -> None:
        table = document.add_table(rows=1, cols=4)
        table.style = "Light Shading Accent 1"
        headers = ["Records", "Methods", "Selected model", "Holdout macro F1"]
        values = [str(payload["records"]), "12", payload["best_model"], f"{payload['f1_macro']:.3f}"]
        for cell, text in zip(table.rows[0].cells, headers):
            cell.text = text
            self._shade(cell, "D9EAF7")
            for run in cell.paragraphs[0].runs: self._font(run, "Calibri", 9, "0B2545", bold=True)
        row = table.add_row().cells
        for cell, text in zip(row, values):
            cell.text = text
            for run in cell.paragraphs[0].runs: self._font(run, "Calibri", 10, "17202A", bold=True)

    def _add_recommendations_table(self, document: Document, payload: dict[str, Any]) -> None:
        document.add_heading("Selected Imputation Strategies", level=1)
        table = document.add_table(rows=1, cols=3)
        table.style = "Light Shading Accent 1"
        for cell, value in zip(table.rows[0].cells, ["Feature", "Recommended method", "Benchmark RMSE"]):
            cell.text = value; self._shade(cell, "D9EAF7")
            for run in cell.paragraphs[0].runs: self._font(run, "Calibri", 9, "0B2545", bold=True)
        for item in payload["recommendations"][:10]:
            cells = table.add_row().cells
            for cell, value in zip(cells, [item["feature"], item["method"], f"{item['rmse']:.3f}"]):
                cell.text = str(value)
                for run in cell.paragraphs[0].runs: self._font(run, "Calibri", 9, "17202A")

    def _image(self, document: Document, filename: str, caption: str) -> None:
        source = self.figures_dir / filename
        if not source.exists():
            return
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(source), width=Inches(6.2))
        caption_p = document.add_paragraph(caption)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption_p.runs: self._font(run, "Calibri", 8.5, "566573", italic=True)

    @staticmethod
    def _font(run: Any, name: str, size: float, color: str, bold: bool = False, italic: bool = False) -> None:
        run.font.name = name; run.font.size = Pt(size); run.bold = bold; run.italic = italic
        run.font.color.rgb = RGBColor.from_string(color)
        run._element.rPr.rFonts.set(qn("w:ascii"), name); run._element.rPr.rFonts.set(qn("w:hAnsi"), name)

    @staticmethod
    def _shade(cell: Any, fill: str) -> None:
        properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd"); shading.set(qn("w:fill"), fill); properties.append(shading)
